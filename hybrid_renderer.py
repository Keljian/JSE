"""
hybrid_renderer.py — Workstream B of the doc-gen pipeline.

Renders a model-authored Markdown body into a clean, consistently-styled .docx.
"Hybrid" = an optional styled base template supplies letterhead / contact / fonts,
and the model-authored body (which decides its own sections, order and emphasis)
is laid out underneath. This replaces the brittle placeholder-surgery in
application_doc_builder.py, giving the model Claude-like structural freedom while
the renderer guarantees the formatting.

Markdown the model is expected to emit:
    # Name / title          -> document title (centered) — skipped if base has a header
    ## SECTION HEADING       -> bold, uppercased, bottom rule
    ### Role | Company | Dates  -> bold subheading
    * bullet  /  - bullet    -> bulleted list item (supports **bold**)
    **bold** inline anywhere
    (blank line)             -> spacing
    plain text               -> body paragraph
"""
from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _bottom_border(paragraph):
    try:
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "888888")
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception:
        pass


def _add_inline(paragraph, text):
    """Render the small Markdown subset used by generated documents."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*))")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        run = paragraph.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        if token.startswith("**"):
            run.bold = True
        else:
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _role_heading_with_date(line):
    """Give undated role headings an explicit, honest chronology label."""
    text = line.strip()
    has_date = bool(re.search(
        r"\b(?:19|20)\d{2}\b|\bpresent\b|\bcurrent\b|\bprior experience\b",
        text,
        re.I,
    ))
    return text if has_date else f"{text} | Prior experience"


def _contact_header(doc, info):
    info = info or {}
    name = f"{info.get('first_name', '')} {info.get('last_name', '')}".strip() or "Curriculum Vitae"
    name_p = doc.add_paragraph()
    r = name_p.add_run(name)
    r.bold = True
    r.font.size = Pt(20)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)

    bits = [info.get("phone"), info.get("email"), info.get("linkedin")]
    contact = "  |  ".join(b for b in bits if b)
    if contact:
        c = doc.add_paragraph()
        cr = c.add_run(contact)
        cr.font.size = Pt(10)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)


def render_markdown_to_docx(markdown_text, output_path, personal_info=None,
                            base_template_path=None, add_header=True):
    if base_template_path and Path(base_template_path).exists():
        doc = docx.Document(str(base_template_path))   # letterhead/styles come from here
        header_present = True
    else:
        doc = docx.Document()
        for section in doc.sections:
            section.left_margin = section.right_margin = Inches(0.9)
            section.top_margin = section.bottom_margin = Inches(0.8)
        header_present = False

    if add_header and not header_present:
        _contact_header(doc, personal_info)

    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        line = line.strip()

        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(12)
            _bottom_border(p)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            _add_inline(p, _role_heading_with_date(line[4:]))
            for run in p.runs:
                run.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith(("* ", "- ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            _add_inline(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.1

    doc.save(str(output_path))
    return str(output_path)


_SIGNOFFS = ("kind regards", "best regards", "warm regards", "regards", "sincerely",
             "yours faithfully", "yours sincerely", "many thanks", "thank you")


def render_cover_letter_to_docx(content, output_path, personal_info=None, base_template_path=None):
    """Render a cover letter with real letter structure: styled sender header,
    date, recipient block, Re: line, greeting, justified body, signoff."""
    info = personal_info or {}
    if base_template_path and Path(base_template_path).exists():
        doc = docx.Document(str(base_template_path))
    else:
        doc = docx.Document()
        for section in doc.sections:
            section.left_margin = section.right_margin = Inches(1.0)
            section.top_margin = section.bottom_margin = Inches(0.65)

    # A compact, readable business-letter default. Together with the authoring
    # word limit this keeps generated letters to one page.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # No centered resume-style banner on a letter — it should open with the
    # date / recipient block like a proper business letter.
    name = f"{info.get('first_name', '')} {info.get('last_name', '')}".strip().lower()
    contact_pat = re.compile(r"(@|linkedin|\b04\d\d|\b\+?61|\d{4}\s?\d{3}\s?\d{3})", re.I)

    lines = [l.rstrip() for l in str(content or "").splitlines()]
    # drop the letter's own leading name/contact lines (styled header replaces them)
    i = 0
    while i < len(lines) and (not lines[i].strip()
                              or lines[i].strip().lower() == name
                              or contact_pat.search(lines[i])):
        # stop skipping once we hit a date or "Dear"/"Re:" — that's real content
        if re.search(r"\b(19|20)\d\d\b", lines[i]) or lines[i].strip().lower().startswith(("dear", "re:")):
            break
        i += 1

    body_started = False
    for line in lines[i:]:
        s = line.strip()
        if not s:
            continue
        low = s.lower()

        if re.fullmatch(r"(re:|subject:).*", low):
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif low.startswith("dear "):
            p = doc.add_paragraph()
            _add_inline(p, s)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            body_started = True
        elif low.rstrip(",") in _SIGNOFFS:
            p = doc.add_paragraph()
            _add_inline(p, s)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
        elif not body_started:
            # date / recipient block — tight, no justification
            p = doc.add_paragraph()
            _add_inline(p, s)
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            _add_inline(p, s)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)

    doc.save(str(output_path))
    return str(output_path)


# --------------------------------------------------------------------------- #
# Demo helper: turn a plain-text resume from the corpus into Markdown so we can
# show the renderer on REAL content (no invention) without an LLM in the loop.
# --------------------------------------------------------------------------- #
def plaintext_resume_to_markdown(text: str) -> str:
    out = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            out.append("")
            continue
        upper = line.upper()
        # all-caps short line => section heading
        if line == upper and 1 <= len(line.split()) <= 5 and len(line) <= 40:
            out.append(f"## {line.title()}")
        # "Company | ... | Dates" or a line ending in a date range => role subheading
        elif ("|" in line and re.search(r"\d{4}", line)) or re.search(r"\b(19|20)\d{2}\s*[–-]\s*((19|20)\d{2}|present)\b", line, re.I):
            out.append(f"### {line}")
        elif line.startswith(("•", "·", "-", "*", "‣", "▪")):
            out.append(f"* {line.lstrip('•·-*‣▪ ').strip()}")
        else:
            out.append(line)
    return "\n".join(out)


if __name__ == "__main__":
    import sys, sqlite3, os
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    data = os.environ.get("JSE_DATA_DIR") or os.path.join(os.environ["APPDATA"], "JSE")
    conn = sqlite3.connect(os.path.join(data, "job_applications.db"))
    conn.row_factory = sqlite3.Row
    # render a real resume + a real cover letter from the corpus to show formatting
    want = sys.argv[1] if len(sys.argv) > 1 else ""
    if not want:
        print("usage: python hybrid_renderer.py <context-document-filename>")
        sys.exit(2)
    row = conn.execute("SELECT filename, doc_type, text FROM context_documents WHERE filename = ?", (want,)).fetchone()
    if not row:
        print(f"not found: {want}")
        sys.exit(1)
    info = {"first_name": "Candidate", "last_name": "", "phone": "",
            "email": "", "linkedin": ""}
    out = Path("applications") / "_render_demo.docx"
    if row["doc_type"] == "resume":
        md = plaintext_resume_to_markdown(row["text"])
    else:
        md = row["text"]
    render_markdown_to_docx(md, out, personal_info=info)
    print(f"Rendered {row['filename']} ({row['doc_type']}) -> {out}")
