"""
context_library.py — Workstream A of the doc-gen context pipeline.

Ingests the candidate's corpus of prior resumes / cover letters / KSC responses, caches
them locally, and retrieves the most relevant evidence for a target job so the
in-app document generator gets the same rich context a candidate may otherwise
feed into an external writing project by hand.

Dependency-light on purpose: extraction uses python-docx / pdfplumber / antiword
(all already present); retrieval is a pure-Python TF-IDF cosine — no sklearn /
embeddings needed for v1.

CLI:
    python context_library.py ingest            # build/refresh the library
    python context_library.py stats             # show what's indexed
    python context_library.py demo <job_id>     # print assembled context for a job
"""
from __future__ import annotations

import hashlib
import math
import os
import re
import shutil
import sqlite3
import subprocess
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

# --- Path resolution (works whether launched by Electron or from C:\JSE) ---
APP_ROOT = Path(os.environ.get("JSE_APP_ROOT") or Path(__file__).resolve().parent)
DATA_DIR = Path(os.environ.get("JSE_DATA_DIR") or APP_ROOT / "settings")
DB_PATH = DATA_DIR / "job_applications.db"
CORPUS_CACHE = DATA_DIR / "context_corpus"
DEFAULT_SOURCE = str(APP_ROOT / "older_applications")
DATA_DIR.mkdir(parents=True, exist_ok=True)

TEXT_EXTS = {".docx", ".doc", ".pdf", ".txt", ".md"}

STOPWORDS = set("""
a an the and or but if then else for to of in on at by with from as is are was were be been being this that these those
it its their his her your our my we you they i he she them us me will would shall should can could may might must not no
your role roles position will be have has had do does did about into over under more most other such only own same so than
job application applicant candidate company team work working experience skills ability please apply role responsibilities
""".split())


# --------------------------------------------------------------------------- #
# DB
# --------------------------------------------------------------------------- #
def _connect():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def ensure_schema(conn):
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS context_documents (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            source_path   TEXT UNIQUE,
            cache_path    TEXT,
            filename      TEXT,
            doc_type      TEXT,
            role_family   TEXT,
            char_len      INTEGER,
            text          TEXT,
            fingerprint   TEXT,
            source_mtime  TEXT,
            indexed_at    TEXT
        )
        """
    )
    conn.commit()


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #
def _extract_docx(path: Path) -> str:
    import docx
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(t for t in parts if t and t.strip())


def _extract_pdf(path: Path) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _extract_doc(path: Path) -> str:
    antiword = shutil.which("antiword")
    if antiword:
        try:
            res = subprocess.run(
                [antiword, str(path)], capture_output=True, timeout=60
            )
            text = res.stdout.decode("utf-8", errors="replace").strip()
            if text:
                return text
        except Exception:
            pass
    # Fallback: Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return text or ""
    except Exception:
        return ""


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return _extract_docx(path)
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".doc":
            return _extract_doc(path)
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f""  # extraction failed; caller logs
    return ""


# --------------------------------------------------------------------------- #
# Classification
# --------------------------------------------------------------------------- #
def classify(filename: str, text: str) -> str:
    name = (filename or "").lower()
    head = (text or "")[:1200].lower()

    if name.startswith("~$"):
        return "other"
    # Employer-authored docs are not candidate evidence.
    if any(k in name for k in ("briefing pack", "candidate brief", "briefing", "candidate pack")):
        return "position_description"
    if "capability statement" in name or "capability statement" in head:
        return "capability_statement"
    # an explicit "cover letter" filename wins even if the body cites selection criteria
    if any(k in name for k in ("cover letter", "cover_letter", "cover", "application letter",
                               "app letter", "covering letter")):
        return "cover_letter"
    if "key selection criteria" in name or "selection criteria" in name or re.search(r"\bksc\b", name):
        return "ksc_response"
    if "position description" in name or re.match(r"^(pd|fed|pr)\d", name) or " - position" in name:
        return "position_description"
    if "letter" in name:
        return "cover_letter"
    # content signals
    if head.startswith("dear ") or "\ndear " in head[:400]:
        if any(s in (text or "").lower() for s in ("sincerely", "kind regards", "yours faithfully", "regards,")):
            return "cover_letter"
    if any(h in (text or "").upper() for h in (
        "PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY", "CAREER EXPERIENCE",
        "KEY SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES", "PROFESSIONAL PROFILE",
    )):
        return "resume"
    if any(k in name for k in ("resume", "cv")):
        return "resume"
    return "other"


ROLE_KEYWORDS = {
    "it_management": ["it manager", "group it", "head of it", "it operations", "service delivery", "infrastructure manager"],
    "infrastructure": ["infrastructure", "cloud", "network", "systems administrator", "devops"],
    "business_analyst": ["business analyst", "systems analyst", "business systems", "ba "],
    "project_delivery": ["project manager", "delivery", "program", "pmo", "portfolio"],
    "engineering": ["engineer", "electronics", "embedded", "pcb", "altium", "hardware", "mechanical"],
    "transformation": ["transformation", "enablement", "digital", "improvement", "change"],
}


def detect_role_family(filename: str, text: str) -> str:
    blob = f"{filename}\n{(text or '')[:1500]}".lower()
    best, best_n = "general", 0
    for family, kws in ROLE_KEYWORDS.items():
        n = sum(blob.count(k) for k in kws)
        if n > best_n:
            best, best_n = family, n
    return best


# --------------------------------------------------------------------------- #
# Ingestion
# --------------------------------------------------------------------------- #
def _fingerprint(path: Path) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def ingest(source: str = DEFAULT_SOURCE, log=print):
    source_root = Path(source)
    if not source_root.is_dir():
        log(f"ERROR: corpus source not reachable: {source_root}")
        return {"error": "source_unreachable"}

    CORPUS_CACHE.mkdir(parents=True, exist_ok=True)
    conn = _connect()
    ensure_schema(conn)

    files = [p for p in source_root.rglob("*") if p.is_file() and p.suffix.lower() in TEXT_EXTS]
    log(f"Found {len(files)} candidate documents under {source_root}")

    stats = Counter()
    for i, path in enumerate(files, 1):
        try:
            rel = path.relative_to(source_root)
        except ValueError:
            rel = Path(path.name)
        source_path = str(path)
        try:
            fp = _fingerprint(path)
        except Exception as e:
            log(f"  [{i}/{len(files)}] SKIP (unreadable) {rel}: {e}")
            stats["unreadable"] += 1
            continue

        existing = conn.execute(
            "SELECT fingerprint FROM context_documents WHERE source_path = ?", (source_path,)
        ).fetchone()
        if existing and existing["fingerprint"] == fp:
            stats["unchanged"] += 1
            continue

        # cache locally (flatten under cache, preserving subfolder as prefix)
        cache_name = re.sub(r"[^A-Za-z0-9._-]+", "_", str(rel))
        cache_path = CORPUS_CACHE / cache_name
        try:
            shutil.copy2(path, cache_path)
        except Exception:
            cache_path = path  # fall back to source if copy fails

        text = _extract_text(Path(cache_path))
        if not text or len(text.strip()) < 40:
            log(f"  [{i}/{len(files)}] thin/empty extract: {rel} ({len(text)} chars)")
            stats["empty"] += 1
            # still record it so we know it exists
        doc_type = classify(path.name, text)
        role = detect_role_family(path.name, text)
        conn.execute(
            """
            INSERT INTO context_documents
                (source_path, cache_path, filename, doc_type, role_family, char_len, text, fingerprint, source_mtime, indexed_at)
            VALUES (?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(source_path) DO UPDATE SET
                cache_path=excluded.cache_path, filename=excluded.filename, doc_type=excluded.doc_type,
                role_family=excluded.role_family, char_len=excluded.char_len, text=excluded.text,
                fingerprint=excluded.fingerprint, source_mtime=excluded.source_mtime, indexed_at=excluded.indexed_at
            """,
            (
                source_path, str(cache_path), path.name, doc_type, role, len(text), text, fp,
                datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
        stats[doc_type] += 1
        if i % 25 == 0:
            conn.commit()
            log(f"  …{i}/{len(files)} processed")
    conn.commit()
    conn.close()
    log(f"Ingest complete: {dict(stats)}")
    return dict(stats)


# --------------------------------------------------------------------------- #
# Retrieval (TF-IDF cosine, pure python)
# --------------------------------------------------------------------------- #
_TOKEN_RE = re.compile(r"[a-z][a-z0-9+#.]{2,}")


def _tokenize(text: str):
    return [t for t in _TOKEN_RE.findall((text or "").lower()) if t not in STOPWORDS]


def _load_index(conn):
    rows = conn.execute(
        "SELECT id, filename, doc_type, role_family, char_len, text, source_mtime "
        "FROM context_documents WHERE char_len >= 40 AND filename NOT LIKE '~$%'"
    ).fetchall()
    docs = []
    df = Counter()
    for r in rows:
        toks = _tokenize(r["text"])
        if not toks:
            continue
        tf = Counter(toks)
        docs.append({"id": r["id"], "filename": r["filename"], "doc_type": r["doc_type"],
                     "role_family": r["role_family"], "char_len": r["char_len"], "tf": tf,
                     "text": r["text"], "mtime": r["source_mtime"] or ""})
        for term in tf:
            df[term] += 1
    n = len(docs)
    idf = {term: math.log((n + 1) / (d + 1)) + 1.0 for term, d in df.items()}
    for doc in docs:
        vec = {t: f * idf.get(t, 0.0) for t, f in doc["tf"].items()}
        norm = math.sqrt(sum(v * v for v in vec.values())) or 1.0
        doc["vec"] = {t: v / norm for t, v in vec.items()}
    return docs, idf


def _score(query_vec, doc):
    return sum(query_vec.get(t, 0.0) * v for t, v in doc["vec"].items())


def _recency_factor(mtime: str) -> float:
    try:
        year = int((mtime or "")[:4])
    except (ValueError, TypeError):
        return 1.0
    if year >= 2024:
        return 1.15
    if year >= 2022:
        return 1.05
    if year >= 2019:
        return 0.95
    return 0.8  # pre-2019 — likely stale voice/role


def retrieve(conn, job_text: str, idf=None, docs=None, per_type=None):
    if docs is None:
        docs, idf = _load_index(conn)
    per_type = per_type or {"resume": 2, "cover_letter": 3, "ksc_response": 2, "capability_statement": 1}
    job_family = detect_role_family("", job_text)
    qtf = Counter(_tokenize(job_text))
    qvec = {t: f * idf.get(t, 0.0) for t, f in qtf.items()}
    qnorm = math.sqrt(sum(v * v for v in qvec.values())) or 1.0
    qvec = {t: v / qnorm for t, v in qvec.items()}

    adjusted = []
    for d in docs:
        base = _score(qvec, d)
        if base <= 0:
            continue
        fam_bonus = 1.3 if d["role_family"] == job_family else 1.0
        adjusted.append((base * fam_bonus * _recency_factor(d.get("mtime")), base, d))
    adjusted.sort(key=lambda x: x[0], reverse=True)

    picked = {k: [] for k in per_type}
    for adj, base, d in adjusted:
        dt = d["doc_type"] if d["doc_type"] in picked else None
        if dt and len(picked[dt]) < per_type[dt]:
            picked[dt].append((round(adj, 4), d))
    return picked, job_family


# --------------------------------------------------------------------------- #
# Assembly — the "Claude project replica"
# --------------------------------------------------------------------------- #
TYPE_BUDGET = {"resume": 6000, "cover_letter": 3500, "ksc_response": 4000, "capability_statement": 2500}
TYPE_LABEL = {
    "resume": "MOST-RELEVANT PRIOR RESUMES",
    "cover_letter": "PRIOR COVER LETTERS (voice & tone exemplars)",
    "ksc_response": "KEY SELECTION CRITERIA / EVIDENCE BANK",
    "capability_statement": "CAPABILITY STATEMENTS",
}


def assemble_context(conn, job_text: str):
    picked, job_family = retrieve(conn, job_text)
    blocks = ["=" * 70,
              f"CANDIDATE EVIDENCE LIBRARY (auto-retrieved; target role family: {job_family})",
              "=" * 70]
    selected = []
    for dt in ("resume", "cover_letter", "ksc_response", "capability_statement"):
        items = picked.get(dt) or []
        if not items:
            continue
        blocks.append(f"\n##### {TYPE_LABEL[dt]} #####")
        for score, d in items:
            budget = TYPE_BUDGET[dt]
            body = (d["text"] or "").strip()
            if len(body) > budget:
                body = body[:budget] + "\n…[truncated]…"
            blocks.append(f"\n--- {d['filename']}  (relevance {score}, family={d['role_family']}) ---\n{body}")
            selected.append({"filename": d["filename"], "doc_type": dt, "score": score})
    return "\n".join(blocks), selected


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def _job_text(conn, job_id):
    row = conn.execute(
        "SELECT title, company, location, description FROM jobs WHERE id = ?", (job_id,)
    ).fetchone()
    if not row:
        return None, None
    label = f"{row['title']} @ {row['company']} ({row['location']})"
    text = f"{row['title']}\n{row['company']}\n{row['description'] or ''}"
    return label, text


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    cmd = sys.argv[1] if len(sys.argv) > 1 else "stats"
    conn = _connect()
    ensure_schema(conn)

    if cmd == "ingest":
        source = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_SOURCE
        ingest(source)
    elif cmd == "reclassify":
        # re-tag already-extracted docs without re-reading files; drop ~$ temp files
        removed = conn.execute("DELETE FROM context_documents WHERE filename LIKE '~$%'").rowcount
        rows = conn.execute("SELECT id, filename, text FROM context_documents").fetchall()
        changed = 0
        for r in rows:
            dt = classify(r["filename"], r["text"])
            rf = detect_role_family(r["filename"], r["text"])
            cur = conn.execute(
                "UPDATE context_documents SET doc_type=?, role_family=? WHERE id=? AND (doc_type IS NOT ? OR role_family IS NOT ?)",
                (dt, rf, r["id"], dt, rf),
            )
            changed += cur.rowcount
        conn.commit()
        print(f"Reclassified {changed} docs; removed {removed} temp (~$) files.")
    elif cmd == "stats":
        rows = conn.execute(
            "SELECT doc_type, COUNT(*), SUM(char_len) FROM context_documents GROUP BY doc_type ORDER BY 2 DESC"
        ).fetchall()
        total = conn.execute("SELECT COUNT(*) FROM context_documents").fetchone()[0]
        print(f"Indexed documents: {total}")
        for r in rows:
            print(f"  {r[0]:24} {r[1]:>4}   {(r[2] or 0):>9,} chars")
    elif cmd == "demo":
        job_id = int(sys.argv[2])
        label, text = _job_text(conn, job_id)
        if not text:
            print(f"Job {job_id} not found.")
            return
        print(f"TARGET JOB: {label}\n")
        block, selected = assemble_context(conn, text)
        print(block)
        print("\n" + "=" * 70)
        print("SELECTED EVIDENCE:")
        for s in selected:
            print(f"  [{s['doc_type']:22}] {s['score']:.3f}  {s['filename']}")
    else:
        print(__doc__)
    conn.close()


if __name__ == "__main__":
    main()
