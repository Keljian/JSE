"""Render AI-generated application content into DOCX templates.

The renderer keeps formatting from user-supplied templates where possible,
replaces known placeholders, and removes unresolved placeholder paragraphs so
generated documents are usable without exposing template internals.
"""
from __future__ import annotations

import copy
import json
import re
from pathlib import Path
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


RESUME_PLACEHOLDERS = {
    "{{professional_profile}}",
    "{{core_skills}}",
    "{{professional_experience}}",
    "{{targeted_experience_bullets}}",
}

COVER_PLACEHOLDERS = {
    "{{cover_letter_subject}}",
    "{{cover_letter_greeting}}",
    "{{cover_letter_opening}}",
    "{{cover_letter_body}}",
    "{{cover_letter_value_proposition}}",
    "{{cover_letter_closing}}",
    "{{cover_letter_signoff}}",
}


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "", value or "application").strip()
    return re.sub(r"\s+", "_", cleaned)[:90] or "application"


def _as_list(value):
    if isinstance(value, list):
        return [item for item in value if item]
    if isinstance(value, str) and value.strip():
        return [line.strip(" -\t") for line in value.splitlines() if line.strip(" -\t")]
    return []


def _paragraphs_in_document(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _paragraph_index(document, paragraph):
    for index, item in enumerate(document.paragraphs):
        if item._p is paragraph._p:
            return index
    return None


def _remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def _is_resume_heading(text):
    stripped = str(text or "").strip()
    return bool(stripped) and stripped.upper() == stripped and len(stripped.split()) <= 5


def _replace_section_after_heading(document, headings, writer):
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().upper() not in headings:
            continue
        end = len(paragraphs)
        for next_index in range(index + 1, len(paragraphs)):
            if _is_resume_heading(paragraphs[next_index].text):
                end = next_index
                break
        anchor = paragraph
        for old in list(paragraphs[index + 1:end]):
            _remove_paragraph(old)
        target = _insert_after(anchor, anchor)
        writer(target)
        return True
    return False


def _insert_after(paragraph, source=None):
    new_element = copy.deepcopy((source or paragraph)._p)
    paragraph._p.addnext(new_element)
    new_para = docx.text.paragraph.Paragraph(new_element, paragraph._parent)
    _clear_paragraph(new_para)
    return new_para


def _clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _set_text(paragraph, text, *, bold=None, italic=None):
    _clear_paragraph(paragraph)
    run = paragraph.add_run(str(text or ""))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def _copy_run_style(source_paragraph, target_paragraph):
    if not source_paragraph.runs or not target_paragraph.runs:
        return
    source = source_paragraph.runs[0]
    target = target_paragraph.runs[0]
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size


def _replace_paragraph_with_lines(paragraph, lines):
    lines = [line for line in lines if str(line or "").strip()]
    if not lines:
        _set_text(paragraph, "")
        return paragraph
    _set_text(paragraph, lines[0])
    _copy_run_style(paragraph, paragraph)
    current = paragraph
    for line in lines[1:]:
        current = _insert_after(current, paragraph)
        _set_text(current, line)
    return current


def _add_bullet_after(anchor, text, source=None):
    paragraph = _insert_after(anchor, source or anchor)
    _set_text(paragraph, f"• {text}")
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.first_line_indent = Pt(-8)
    return paragraph


def _flatten_skills(core_skills):
    if isinstance(core_skills, dict):
        values = []
        for group, items in core_skills.items():
            values.append(str(group))
            values.extend([f"- {item}" for item in _as_list(items)])
        return values
    return [f"- {item}" for item in _as_list(core_skills)]


def _replace_skills(paragraph, skills):
    values = _flatten_skills(skills)
    if not values:
        values = ["Relevant skills to be tailored for this role."]
    first = str(values[0])
    if first.startswith("- "):
        _set_text(paragraph, f"• {first[2:]}")
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.first_line_indent = Pt(-8)
    else:
        _set_text(paragraph, first, bold=True)
    current = paragraph
    for value in values[1:]:
        if str(value).startswith("- "):
            current = _add_bullet_after(current, str(value)[2:], paragraph)
        else:
            current = _insert_after(current, paragraph)
            _set_text(current, value, bold=True)
    return current


def _replace_experience(paragraph, experience):
    roles = _as_list(experience)
    if not roles:
        _set_text(paragraph, "Relevant experience to be tailored for this role.")
        return paragraph

    current = paragraph
    first = True
    for role in roles:
        if not isinstance(role, dict):
            if first:
                _set_text(current, str(role))
                first = False
            else:
                current = _insert_after(current, paragraph)
                _set_text(current, str(role))
            continue

        if not first:
            current = _insert_after(current, paragraph)
        first = False

        company = role.get("company") or ""
        title = role.get("title") or role.get("role_title") or ""
        dates = role.get("dates") or ""
        header = " | ".join(part for part in [company, dates] if part)
        _set_text(current, header or title, bold=True)
        current.paragraph_format.space_before = Pt(6)

        if title and header != title:
            current = _insert_after(current, paragraph)
            _set_text(current, title, bold=True)

        summary = role.get("summary") or role.get("context") or ""
        if summary:
            current = _insert_after(current, paragraph)
            _set_text(current, summary, italic=True)
            current.paragraph_format.space_after = Pt(4)

        achievements = role.get("achievements") or role.get("bullets") or []
        if achievements:
            current = _insert_after(current, paragraph)
            _set_text(current, "Key Achievements:")
            current.paragraph_format.space_before = Pt(4)
            for bullet in _as_list(achievements)[:8]:
                current = _add_bullet_after(current, bullet, paragraph)

    return current


def _replace_cover_body(paragraph, value):
    if isinstance(value, list):
        return _replace_paragraph_with_lines(paragraph, value)
    return _replace_paragraph_with_lines(paragraph, [part.strip() for part in str(value or "").split("\n\n")])


def _find_cover_value(content, token):
    cover = content.get("cover_letter") if isinstance(content.get("cover_letter"), dict) else {}
    key = token.strip("{}")
    return content.get(key) or cover.get(key.replace("cover_letter_", "")) or cover.get(key)


def render_resume_template(template_path, output_path, content):
    document = docx.Document(str(template_path))
    _replace_section_after_heading(
        document,
        {"PROFESSIONAL PROFILE", "PROFILE", "PROFESSIONAL SUMMARY"},
        lambda paragraph: _replace_paragraph_with_lines(paragraph, [content.get("professional_profile", "")]),
    )
    _replace_section_after_heading(
        document,
        {"KEY SKILLS", "TECHNICAL COMPETENCIES", "CORE COMPETENCIES", "TECHNICAL SKILLS"},
        lambda paragraph: _replace_skills(paragraph, content.get("core_skills") or []),
    )
    _replace_section_after_heading(
        document,
        {"PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY", "CAREER EXPERIENCE", "EXPERIENCE"},
        lambda paragraph: _replace_experience(paragraph, content.get("professional_experience") or content.get("experience") or []),
    )
    for paragraph in list(_paragraphs_in_document(document)):
        text = paragraph.text.strip()
        if text == "{{professional_profile}}":
            _replace_paragraph_with_lines(paragraph, [content.get("professional_profile", "")])
        elif text == "{{core_skills}}":
            _replace_skills(paragraph, content.get("core_skills") or [])
        elif text in {"{{professional_experience}}", "{{targeted_experience_bullets}}"}:
            _replace_experience(paragraph, content.get("professional_experience") or content.get("experience") or [])

    for paragraph in list(_paragraphs_in_document(document)):
        if "{{" in paragraph.text and "}}" in paragraph.text:
            _remove_paragraph(paragraph)

    document.save(str(output_path))
    return str(output_path)


def render_cover_letter_template(template_path, output_path, content):
    document = docx.Document(str(template_path))
    cover_signoff = _find_cover_value(content, "{{cover_letter_signoff}}")
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(list(paragraphs)):
        text = paragraph.text.strip()
        if re.fullmatch(r"\d{1,2}/\d{1,2}/\d{4}", text):
            _replace_paragraph_with_lines(paragraph, [datetime.now().strftime("%d/%m/%Y")])
        if text.lower().startswith("dear "):
            value = _find_cover_value(content, "{{cover_letter_greeting}}")
            if value:
                _replace_paragraph_with_lines(paragraph, [value])
        if text == "{{cover_letter_closing}}":
            for old in list(document.paragraphs[index + 1:]):
                old_text = old.text.strip().lower()
                if old_text in {"sincerely,", "sincerely", "kind regards,", "kind regards", "regards,", "regards"}:
                    break
                if old.text.strip():
                    _remove_paragraph(old)
            break

    if cover_signoff:
        for index, paragraph in enumerate(list(document.paragraphs)):
            old_text = paragraph.text.strip().lower()
            if old_text in {"sincerely,", "sincerely", "kind regards,", "kind regards", "regards,", "regards"}:
                lines = [line.strip() for line in str(cover_signoff).splitlines() if line.strip()]
                _replace_paragraph_with_lines(paragraph, lines)
                for old in list(document.paragraphs[index + len(lines):]):
                    old_text = old.text.strip()
                    if old_text and old_text.lower() not in {"sincerely,", "sincerely", "kind regards,", "kind regards", "regards,", "regards"}:
                        _remove_paragraph(old)
                        break
                break

    for paragraph in list(_paragraphs_in_document(document)):
        text = paragraph.text.strip()
        if text in COVER_PLACEHOLDERS:
            value = _find_cover_value(content, text)
            if text == "{{cover_letter_body}}":
                _replace_cover_body(paragraph, value)
            else:
                _replace_paragraph_with_lines(paragraph, [value or ""])
    for paragraph in list(_paragraphs_in_document(document)):
        if "{{" in paragraph.text and "}}" in paragraph.text:
            _remove_paragraph(paragraph)
    document.save(str(output_path))
    return str(output_path)


def write_generation_json(output_path, content):
    Path(output_path).write_text(json.dumps(content, indent=2, ensure_ascii=False), encoding="utf-8")
    return str(output_path)
