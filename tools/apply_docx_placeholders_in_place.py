"""Replace text markers in DOCX templates with placeholder tokens in place."""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import site

sys.path.append(site.getusersitepackages())

from docx import Document


ROOT = Path("Application templates")
BACKUP_ROOT = ROOT / "_originals_before_placeholders"


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def normalize(text: str) -> str:
    return " ".join(text.strip().lower().replace(":", "").split())


def restore_originals() -> tuple[list[Path], list[str]]:
    restored: list[Path] = []
    skipped: list[str] = []
    if not BACKUP_ROOT.exists():
        return restored, skipped
    for backup in BACKUP_ROOT.rglob("*.docx"):
        relative = backup.relative_to(BACKUP_ROOT)
        target = ROOT / relative
        if target.name.startswith("~$"):
            continue
        target.parent.mkdir(parents=True, exist_ok=True)
        try:
            shutil.copy2(backup, target)
            restored.append(target)
        except PermissionError:
            skipped.append(f"locked {target}")
    return restored, skipped


def classify(path: Path, document: Document) -> str:
    all_text = "\n".join(paragraph.text for paragraph in iter_paragraphs(document)).lower()
    parent = str(path.parent).lower()
    name = path.name.lower()
    first_page_text = all_text[:1800]

    if "cvs" in parent:
        return "resume"
    if "cover letter" in parent:
        return "cover_letter"
    if "key selection criteria" in all_text or "selection criteria" in first_page_text:
        return "selection_criteria"
    if "dear " in first_page_text or "to the hiring manager" in first_page_text or "to whom it may concern" in first_page_text:
        return "cover_letter"
    if "professional profile" in all_text or "resume" in name or "cv" in name:
        return "resume"
    if "key experience" in first_page_text:
        return "selection_criteria"
    return "application_document"


def set_placeholder(paragraph, placeholder: str) -> bool:
    if placeholder in paragraph.text:
        return False
    paragraph.text = placeholder
    if paragraph.runs:
        run = paragraph.runs[0]
        run.italic = True
        run.bold = False
    return True


def next_non_empty(paragraphs, start_index: int):
    for index in range(start_index + 1, len(paragraphs)):
        if paragraphs[index].text.strip():
            return index, paragraphs[index]
    return None, None


def replace_after_heading(paragraphs, headings: set[str], placeholder: str) -> int:
    changed = 0
    for index, paragraph in enumerate(paragraphs):
        if normalize(paragraph.text) in headings:
            _, target = next_non_empty(paragraphs, index)
            if target is not None:
                changed += int(set_placeholder(target, placeholder))
            break
    return changed


def paragraph_after_greeting(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        text = normalize(paragraph.text)
        if text.startswith("dear ") or text in {"to the hiring manager", "to whom it may concern"}:
            return index
    return None


def apply_resume(document: Document) -> int:
    paragraphs = list(iter_paragraphs(document))
    changed = 0
    changed += replace_after_heading(
        paragraphs,
        {"professional profile", "profile", "professional summary", "career summary"},
        "{{professional_profile}}",
    )
    changed += replace_after_heading(
        paragraphs,
        {"career highlights", "selected achievements", "achievements"},
        "{{selected_achievements}}",
    )
    changed += replace_after_heading(
        paragraphs,
        {
            "technical competencies",
            "core competencies",
            "key attributes",
            "skills",
            "technical skills",
            "capabilities",
        },
        "{{core_skills}}",
    )
    changed += replace_after_heading(
        paragraphs,
        {
            "professional experience",
            "employment history",
            "career experience",
            "experience",
            "work experience",
        },
        "{{targeted_experience_bullets}}",
    )
    changed += replace_after_heading(
        paragraphs,
        {"education", "qualifications", "education and qualifications"},
        "{{education_section}}",
    )
    return changed


def apply_cover_letter(document: Document) -> int:
    paragraphs = list(iter_paragraphs(document))
    changed = 0

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        lower = text.lower()
        if lower.startswith("re:") or lower.startswith("subject:"):
            changed += int(set_placeholder(paragraph, "{{cover_letter_subject}}"))
            break

    greeting_index = paragraph_after_greeting(paragraphs)
    if greeting_index is None:
        return changed

    targets = [
        "{{cover_letter_opening}}",
        "{{cover_letter_body}}",
        "{{cover_letter_value_proposition}}",
        "{{cover_letter_closing}}",
    ]
    cursor = greeting_index
    for placeholder in targets:
        next_index, target = next_non_empty(paragraphs, cursor)
        if target is None:
            break
        changed += int(set_placeholder(target, placeholder))
        cursor = next_index
    return changed


def apply_selection_criteria(document: Document) -> int:
    paragraphs = list(iter_paragraphs(document))
    changed = 0
    changed += replace_after_heading(
        paragraphs,
        {"key selection criteria", "selection criteria", "key experience", "addressing key selection criteria"},
        "{{selection_criteria_responses}}",
    )

    response_number = 1
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not text or "{{" in text:
            continue
        if text.endswith(":") and len(text.split()) > 5 and response_number <= 5:
            changed += int(set_placeholder(paragraph, f"{{{{selection_criteria_response_{response_number}}}}}"))
            response_number += 1
    return changed


def process(path: Path) -> str:
    try:
        document = Document(str(path))
    except Exception as exc:
        return f"skipped {path}: {exc}"

    kind = classify(path, document)
    if kind == "resume":
        changed = apply_resume(document)
    elif kind == "cover_letter":
        changed = apply_cover_letter(document)
    elif kind == "selection_criteria":
        changed = apply_selection_criteria(document)
    else:
        changed = apply_resume(document) + apply_cover_letter(document) + apply_selection_criteria(document)

    if changed:
        document.save(str(path))
    return f"{'updated' if changed else 'unchanged'} {path}: {kind}, placeholders changed: {changed}"


def main() -> int:
    restored, restore_skips = restore_originals()
    if restored:
        print(f"Restored {len(restored)} DOCX files from {BACKUP_ROOT}")
    else:
        print("No originals restored; backup folder was not found.")
    for message in restore_skips:
        print(message)

    for path in sorted(ROOT.rglob("*.docx")):
        if path.name.startswith("~$") or BACKUP_ROOT in path.parents:
            continue
        print(process(path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
