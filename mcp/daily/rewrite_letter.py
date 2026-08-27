"""Replace the body of a generated cover letter, keeping JSE's letterhead.

The local model drafts all five letters through JSE's templates. For the two that
matter most, a better body gets written by hand. This swaps it in without
rebuilding the document: the date block, the addressee, the Re: line, the
salutation, the sign-off and every formatting choice in the template survive, and
only the paragraphs between "Dear ..." and the sign-off are replaced.

Working on the generated file rather than producing a new one matters: the path is
already recorded on the job row, so nothing downstream has to be told about it.

Usage:
    python rewrite_letter.py --job-id 40665 --body body.txt [--re "Re: ..."]
    python rewrite_letter.py --path "C:\\JSE\\applications\\X_cover_letter.docx" --body body.txt

body.txt is plain text. Blank lines separate paragraphs. A --backup copy is written
beside the original unless --no-backup is given.
"""

import argparse
import copy
import re
import shutil
import sqlite3
import sys
from pathlib import Path

from docx import Document

DB = Path(r"C:\JSE\settings\job_applications.db")

SIGNOFFS = ("yours sincerely", "yours faithfully", "kind regards", "regards,", "sincerely,", "best regards")


def letter_path_for_job(job_id):
    con = sqlite3.connect(str(DB))
    try:
        row = con.execute("SELECT cover_letter_path, title, company FROM jobs WHERE id = ?", (job_id,)).fetchone()
    finally:
        con.close()
    if not row or not row[0]:
        raise SystemExit(
            f"job {job_id} has no cover_letter_path recorded. Generate documents first "
            "(jse_prepare_applications), then retry."
        )
    path = Path(row[0])
    if not path.exists():
        raise SystemExit(f"recorded path does not exist: {path}")
    return path


def find_bounds(doc):
    """Index of the salutation and of the sign-off."""
    start = end = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if start is None and text.startswith("dear "):
            start = i
        elif start is not None and any(text.startswith(s) for s in SIGNOFFS):
            end = i
            break
    if start is None or end is None:
        raise SystemExit(
            "could not locate the salutation and sign-off in this document, so the body "
            "cannot be replaced safely. Edit it by hand."
        )
    return start, end


def replace_body(path, paragraphs, re_line=None, backup=True):
    if backup:
        shutil.copy2(path, path.with_suffix(".bak.docx"))

    doc = Document(str(path))
    start, end = find_bounds(doc)

    # Keep one of the paragraphs being removed as the formatting donor, so the new
    # text inherits the template's font, size and spacing rather than the theme
    # default. Cloning the XML is the only reliable way to carry run properties.
    donor = None
    for para in doc.paragraphs[start + 1:end]:
        if para.text.strip():
            donor = para
            break
    donor_style = donor.style if donor is not None else doc.paragraphs[start].style
    donor_rpr = None
    if donor is not None and donor.runs:
        donor_rpr = donor.runs[0]._element.get_or_add_rPr()

    signoff = doc.paragraphs[end]
    for para in list(doc.paragraphs[start + 1:end]):
        para._element.getparent().remove(para._element)

    # Blank line after the salutation, matching the template's own spacing.
    signoff.insert_paragraph_before("", donor_style)
    for index, text in enumerate(paragraphs):
        if index:
            signoff.insert_paragraph_before("", donor_style)
        new = signoff.insert_paragraph_before("", donor_style)
        run = new.add_run(text)
        if donor_rpr is not None:
            run._element.insert(0, copy.deepcopy(donor_rpr))
    signoff.insert_paragraph_before("", donor_style)

    if re_line:
        for para in doc.paragraphs[:start]:
            if para.text.strip().lower().startswith("re:"):
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
                if para.runs:
                    para.runs[0].text = re_line
                else:
                    para.add_run(re_line)
                break

    doc.save(str(path))
    return start, end


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--job-id", type=int)
    ap.add_argument("--path")
    ap.add_argument("--body", required=True, help="plain text file, blank lines separate paragraphs")
    ap.add_argument("--re", dest="re_line", default=None)
    ap.add_argument("--no-backup", action="store_true")
    args = ap.parse_args()

    if not args.job_id and not args.path:
        raise SystemExit("give --job-id or --path")
    path = Path(args.path) if args.path else letter_path_for_job(args.job_id)

    raw = Path(args.body).read_text(encoding="utf-8")
    paragraphs = [chunk.strip().replace("\n", " ") for chunk in raw.split("\n\n") if chunk.strip()]
    if not paragraphs:
        raise SystemExit("body file is empty")

    # An em dash in an Australian application reads as machine-written. JSE's own
    # document run on 26 August stripped them; keep that rule enforced here rather
    # than trusting every future draft to remember it.
    paragraphs = [re.sub(r"\s+", " ", p.replace("\u2014", "-").replace("\u2013", "-")).strip() for p in paragraphs]

    start, end = replace_body(path, paragraphs, args.re_line, backup=not args.no_backup)
    print(f"rewrote {len(paragraphs)} paragraphs in {path.name} (replaced body between {start} and {end})")


if __name__ == "__main__":
    main()
